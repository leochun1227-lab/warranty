from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path


OUT = Path("outputs/c4c_sap_variable_dictionary.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None, size=8.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "bottom", "start", "end"):
                el = tc_mar.find(qn(f"w:{side}"))
                if el is None:
                    el = OxmlElement(f"w:{side}")
                    tc_mar.append(el)
                el.set(qn("w:w"), "80" if side in ("top", "bottom") else "100")
                el.set(qn("w:type"), "dxa")


def add_table(doc, title, rows, widths=(1550, 1450, 2550, 2650, 1160)):
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["变量名", "来源", "业务含义", "代码用途", "影响模块"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, color="0B2545", size=8.5)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row):
            set_cell_text(cells[i], v, size=8.1)
    set_table_geometry(table, widths)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.1)
        run = p.add_run(item)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)


def style_doc(doc):
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Inches(11)
    sec.page_height = Inches(8.5)
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 15, "2E74B5"),
        ("Heading 2", 12.5, "2E74B5"),
        ("Heading 3", 11, "1F4D78"),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(8)
        st.paragraph_format.space_after = Pt(5)


C4C_CORE = [
    ("TicketID / ticketID / id", "C4C Ticket", "C4C 工单唯一编号。相当于所有系统关联的主键。", "用于 Firebase key、历史状态比较、Repair Analytics 明细、SAP 回填匹配、导出。", "全局"),
    ("TicketName", "C4C Ticket", "工单标题或客户/问题摘要。", "作为客户/工单显示兜底字段，也参与状态快照签名，变化时触发同步。", "导出、明细"),
    ("TicketStatus", "C4C Ticket", "C4C 工单状态代码，例如 Z9、Y0、Y1、Y2、Y4、Y8、YB、Z1。", "判断 Critical、Approved/Unapproved；和状态映射表一起决定 dashboard 分类。", "KPI、Repair Analytics、趋势"),
    ("TicketStatusText", "C4C Ticket", "C4C 工单状态文本，例如 New Claim、Repair in Progress、Unapproved Claims Closed。", "状态展示、Critical 判断兜底、Unapproved 判断、历史 from/to 文本。", "状态分布、趋势、明细"),
    ("TicketSeverity", "C4C Ticket", "工单严重程度。", "用于识别 Critical 的兜底条件，也纳入 ticket signature 判断变化。", "Critical 逻辑"),
    ("TicketType / TicketTypeText", "C4C Ticket", "索赔类型文本/代码。当前主要区分 In Field Warranty Claims 和 Pre Delivery Warranty Claims。", "claim filter、KPI claim split、Repair Analytics、Employee Workload 筛选。", "筛选、KPI、导出"),
    ("ClaimType / TicketClaimType", "C4C/派生", "索赔类型的兼容字段名。", "当前字段缺失时作为 TicketTypeText 的别名读取。", "筛选、统计"),
    ("CreatedOn / createdOn / CreatedAt", "C4C Ticket", "工单创建日期。", "New Tickets 趋势、日期范围筛选、aging 起点兜底、导出。", "趋势、Employee、导出"),
    ("LastUpdateDateTime / UpdatedOn / LastChangedOn 等", "C4C Ticket", "工单最后更新时间。", "Unapproved 无明确日期时作为判定日期兜底；导出 Last Update。", "Repair Analytics、导出"),
    ("Responded", "C4C Ticket", "C4C 响应标识。", "作为核心字段同步到 Firebase；当前主要保留用于追溯，不直接驱动 dashboard。", "数据留痕"),
    ("ApprovalDate", "C4C Ticket", "C4C 审批日期字段。", "作为核心字段同步留存；当前 approved 主要由 SAP PO + C4C 状态判断。", "数据留痕"),
    ("ApprovalNumber", "C4C Ticket", "C4C 审批编号。", "同步留存，可用于后续审批链路追溯。", "数据留痕"),
    ("AmountIncludingTax", "C4C Ticket", "含税金额，代表当前维修/索赔金额。", "Approved Cost、Avg Cost/Ticket、维修成本区间、exited value、导出金额。", "成本 KPI、Repair Analytics"),
    ("ServiceRequesterEmail", "C4C Ticket", "服务请求人邮箱。", "同步留存，也可作为客户显示兜底。", "明细、追溯"),
    ("CustomerName / AccountName / BuyerName", "C4C Ticket", "客户、账户或买方名称。", "Repair Analytics、Total Critical 明细、导出里的 Customer。", "明细、导出"),
    ("RepairDescription / Subject / Description / TicketDescription", "C4C Ticket", "维修问题描述或工单主题。", "Repair Analytics 明细里的 Repair 字段。", "Repair Analytics"),
]

CONFIG_FIELDS = [
    ("C4C_USERNAME / C4C_PASSWORD", "环境变量/C4C API", "C4C 接口账号和密码。", "用于登录 C4C Ticket queryOdataBatch；文档不记录实际值。", "抓取配置"),
    ("HOST / PATH", "C4C API 配置", "C4C API host 和路径。当前路径为 Ticket queryOdataBatch。", "决定从哪个 C4C endpoint 拉 ticket 和 involved-party role 数据。", "抓取配置"),
    ("API_TOP / API_SKIP_START / API_EXTRA_TAIL_PAGES", "C4C API 分页配置", "C4C 分页大小、起始 skip、额外尾页数量。", "保证分页拉全，尾页用于处理 C4C totalCount 偶尔滞后的情况。", "抓取稳定性"),
    ("SAP_HANA_DSN", "环境变量/SAP HANA", "SAP HANA ODBC 连接串。", "连接 HANA 查询 VBAK/VBAP/LIPS/LIKP；文档不记录实际值。", "SAP 查询配置"),
    ("SAP_CLIENT", "环境变量/SAP", "SAP client，业务上对应 MANDT。", "SQL 中限制 vbak/lips/vbap 的 MANDT，避免跨 client。", "SAP 查询配置"),
    ("SALES_ORG", "环境变量/SAP", "销售组织，对应 SAP VBAK.VKORG。", "SQL 限定销售订单所属销售组织。", "SAP 查询配置"),
    ("FIREBASE_ROOT", "Firebase 配置", "C4C ticket 原始/增强数据写入根节点，当前逻辑默认 c4cTickets_test。", "保存 ticket core、roles、SAP SO 字段。", "数据存储"),
    ("MONITOR_ROOT", "Firebase 配置", "监控/分析输出根节点，当前逻辑默认 ctmTicketStatusMonitorV44。", "保存 history、analytics、employeeDirectory 等。", "Dashboard 数据源"),
]

C4C_DEALER_EMP = [
    ("DealerID", "C4C Ticket", "经销商编号。", "经销商识别、SAP/明细留存、dealer alias 映射辅助。", "Dealer Workbench、导出"),
    ("DealerName", "C4C Ticket", "经销商名称。", "Top Critical Tickets by Dealer、Dealer Workbench 跳转、明细显示。", "Dealer、Team Dashboard"),
    ("WarrantyHandlingDealerID", "C4C Ticket", "Warranty handling dealer 编号。", "DealerID 为空时辅助识别维修/保修处理经销商。", "Dealer 统计"),
    ("WarrantyHandlingDealerName", "C4C Ticket", "Warranty handling dealer 名称。", "DealerName 为空时兜底。", "Dealer 统计"),
    ("RepairerBusinessNameID", "C4C Ticket", "维修商业务伙伴 ID。", "Repairer/Dealer 兜底字段；帮助定位由哪个维修商处理。", "Dealer、Repair Analytics"),
    ("RepairerEmail", "C4C Ticket", "维修商邮箱。", "同步留存；用于联系维修商或后续明细追溯。", "数据留痕"),
    ("RepairerPhoneNumber", "C4C Ticket", "维修商电话。", "同步留存；用于联系维修商。", "数据留痕"),
    ("RepairerNamePointOfContact", "C4C Ticket", "维修商联系人。", "同步留存；可用于人工 follow-up。", "数据留痕"),
    ("C4CAssignTo / AssignTo / Employee / OwnerName", "C4C Ticket", "票面上的 Assign To/Owner 字段，可能是人，也可能是 Queue Warranty。", "网页端作为 role 40 为空时的兜底；不是首选员工来源。", "Employee Workload"),
    ("Assigned to / Assigned To / AssignedToName / Assignee / AssignedUser / OwnerPartyName", "C4C Ticket", "C4C API 中 Assigned To 的多种拼写。", "同步到 Firebase，并保存 AssignedToRaw 供排查；Z1 新票可显示 Queue Warranty。", "Employee Workload、排查"),
    ("AssignedToRaw", "派生/Firebase", "标准化后的原始 Assigned To 值。", "保留 C4C 票面 Assigned To，帮助解释 role 40 缺失时为什么归到 Queue。", "排查"),
]

C4C_ROLES = [
    ("roles / InvolvedParty*", "C4C involved parties", "C4C 按角色返回的 involved party 集合。不同 role_code 拉回不同当事人。", "抓取脚本按 role 拉取后上传到 Firebase roles 节点。", "Employee、Dealer"),
    ("InvolvedPartyRoleID", "C4C involved parties", "当事人角色代码。", "role 40 代表 Assign To；role 1001 代表 Dealer/经销商。", "员工/经销商归属"),
    ("InvolvedPartyName", "C4C involved parties", "角色对应的人名或组织名。", "role 40 取真实员工 owner；role 1001 取经销商名称。", "Employee Workload、Dealer"),
    ("InvolvedPartyID", "C4C involved parties", "角色对应对象 ID。", "用于角色数据追溯和去重。", "数据留痕"),
    ("InvolvedPartyBusinessPartnerID", "C4C involved parties", "业务伙伴编号。", "用于角色数据追溯和经销商/员工归属辅助。", "数据留痕"),
    ("requested_role_code / requested_role_name / requested_skip", "抓取元数据", "本次 C4C role API 请求的角色和分页信息。", "上传 roles 时用于排查该行来自哪个角色请求。", "排查"),
    ("role 40", "C4C involved parties", "C4C Assign To 员工角色。", "Employee Workload 的权威员工来源；为空且 Z1/New Claim 时归 Queue Warranty。", "Employee Workload"),
    ("role 1001", "C4C involved parties", "经销商/业务伙伴角色。", "DealerName 为空时用 role 1001 识别经销商。", "Dealer 统计"),
]

SAP_FIELDS = [
    ("ERPFreeOrder", "C4C Ticket", "C4C 上的 free order / ERP free order 值，用来反查 SAP 销售订单。", "生成 LookupSalesOrder；所有 SAP HANA 查询入口。", "Parts Delivery、SAP 回填"),
    ("_Lookup1 / _Lookup2", "派生", "从 ERPFreeOrder 生成的 SAP Sales Order 查询候选。", "处理前导零/格式差异，提高 SAP 匹配率。", "SAP 查询"),
    ("MANDT / SAP Client", "SAP HANA", "SAP client，本项目固定使用 800。", "限制只查 MANDT=800，避免跨 client 数据污染。", "SAP 查询"),
    ("VKORG", "SAP VBAK", "销售组织。", "SQL 限定 SALES_ORG，确保销售订单属于正确组织。", "SAP 查询"),
    ("VBAK.VBELN / Sales Order", "SAP VBAK", "SAP 销售订单号。", "表示 C4C ERPFreeOrder 是否已经在 SAP 形成订单。", "Parts Delivery、订单状态"),
    ("VBAK.ERDAT / SO Created Date", "SAP VBAK", "销售订单创建日期。", "回填到 ticket，展示订货日期/订单建立日期。", "Parts Delivery、导出"),
    ("VBAP.POSNR / Sales Order Item", "SAP VBAP", "销售订单行项目号。", "构建 Sales Order Details，每个配件/物料一行。", "Parts Delivery"),
    ("VBAP.MATNR / Material", "SAP VBAP", "物料号。", "显示订了哪个配件/物料。", "Parts Delivery"),
    ("VBAP.ARKTX / Description", "SAP VBAP", "物料描述。", "配件详情说明。", "Parts Delivery"),
    ("VBAP.KWMENG / Order Qty", "SAP VBAP", "订货数量。", "显示每个物料订购数量。", "Parts Delivery"),
    ("VBAP.VRKME / Sales Unit", "SAP VBAP", "销售单位。", "配合 Order Qty 展示单位。", "Parts Delivery"),
    ("VBAP.ABGRU / Rejection Reason", "SAP VBAP", "SAP 销售订单行拒绝原因。", "如果有拒绝原因，该 item 被排除出 Parts Delivery。", "Parts Delivery 过滤"),
    ("Item Rejection Status", "派生/SAP", "由 Rejection Reason 推导：Rejected 或 Not Rejected。", "过滤被 SAP 拒绝的配件行；SAP item 拒绝优先于 C4C 状态。", "Parts Delivery"),
    ("LIPS.VGBEL / LIPS.VGPOS", "SAP LIPS", "交货行对应的原销售订单号和行号。", "关联交货记录到销售订单行。", "Parts Delivery"),
    ("LIPS.VBELN", "SAP LIPS", "交货单号。", "统计 distinct delivery count。", "Parts Delivery"),
    ("LIKP.WADAT_IST", "SAP LIKP", "实际发货/过账日期。", "只有不为空才计入 Delivery Count，表示已经发生实际交货。", "Parts Delivery"),
    ("Delivery Count", "派生/SAP", "某销售订单行已实际交货的交货单数量。", "判断物料是否已经发出/交付。", "Parts Delivery"),
    ("Order Rejection Status", "派生/SAP", "订单层面的拒绝状态。", "Approved PO 判断时必须不是 rejected/partially rejected。", "Approved 判定"),
    ("Issue Status", "派生/SAP", "SAP 订单匹配状态，如 Not Found。", "没有 Sales Order 时标记 Not Found，网页和导出展示。", "Parts Delivery、导出"),
    ("Sales Order Details", "Firebase 派生", "过滤 rejected items 后的 SAP item 明细数组。", "网页 Parts Delivery 使用，包含物料、数量、交货数。", "Parts Delivery"),
]

APPROVAL_FIELDS = [
    ("ERPPurchaseOrder", "C4C/SAP 回写", "ERP 采购订单号。业务上代表采购订单已经生成。", "Approved 条件之一：必须是 7 开头的 10 位 PO。用来看是不是已经订货/采购。", "Approved Tickets、Avg Cost"),
    ("PO Document Date / PODocumentDate / DocumentDate", "SAP/C4C 回写", "采购订单 document date。", "Approved 决策日期优先使用它。", "Repair Analytics、趋势"),
    ("ERPPurchaseOrderFirstSeenDate", "派生/Firebase", "系统第一次看到 ERP PO 的日期。", "PO Document Date 缺失时作为 approved 日期兜底。", "Repair Analytics"),
    ("Approval Decision Date / approvalDecisionDate", "派生", "审批决策日期。", "网页端兼容字段；缺少 PO 日期时兜底。", "Repair Analytics"),
    ("Z1Z8TimeConsumed", "C4C Ticket", "C4C 时间消耗字段，格式如 131 D 17 H 32 M。", "旧 dealer approval rule 曾用 totalMinutes>0 判定 approved；当前团队 dashboard approved 主要改为 PO+状态规则。", "兼容/历史逻辑"),
    ("APPROVED_C4C_STATUS_CODES = Z9/Y0/Y1/Y2/Y4/YB", "代码常量", "这些 C4C 状态被视为可以与有效 PO 一起构成 approved。", "isValidApprovedPo/approvalBucket 判定 approved。", "Repair Analytics、KPI"),
    ("Y8 / Unapproved text", "C4C Ticket/History", "Y8 或文本包含 unapproved 表示未批准关闭。", "isUnapprovedTicket/logIsUnapprovedExit 判定 unapproved。", "Repair Analytics、Employee Approval"),
    ("Unapproved Date", "C4C/派生", "未批准日期。", "Unapproved 决策日期优先使用。", "Repair Analytics"),
]

HISTORY_FIELDS = [
    ("history / logs", "Firebase monitor", "每次状态变化记录。", "计算 entered/exited/moved、unapproved exits、critical trend。", "趋势、KPI"),
    ("cls", "历史事件", "事件类别：enter/exit/move。", "进入 Critical、离开 Critical、状态移动统计。", "Critical Status Trends"),
    ("type / changeType / eventType", "历史事件", "原始事件类型文本。", "保存事件语义和导出追溯。", "日志"),
    ("detectedAt / date", "历史事件", "系统检测到事件的时间/日期。", "趋势日期、approval/unapproved 决策日期、导出。", "趋势、Repair Analytics"),
    ("dataSyncAt", "历史事件", "数据同步时间。", "排查历史事件来自哪次同步。", "排查"),
    ("fromCode / oldCode / fromStatus / oldStatus", "历史事件", "变化前状态代码/文本。", "状态移动分析、历史 tooltip、恢复某日 critical stock。", "趋势、历史"),
    ("toCode / newCode / toStatus / newStatus", "历史事件", "变化后状态代码/文本。", "识别 unapproved exit、exited critical。", "Repair Analytics、趋势"),
    ("amount", "历史事件/派生", "事件上的金额，通常兜底取 ticket amount。", "Repair Analytics unapproved 行金额、exited value。", "成本、导出"),
    ("employee / owner / assignedTo", "历史事件/派生", "事件关联员工。", "当 ticket entry 可匹配时仍优先取 C4C role 40；否则用日志字段兜底。", "Employee Modal"),
    ("agingDays / durationDays", "历史事件/派生", "工单年龄/持续天数。", "Long Duration、Critical Duration Distribution、导出。", "时长分析"),
]

DERIVED_FIELDS = [
    ("isCritical", "派生", "工单当前是否属于 Critical。", "由状态映射/状态文本判断，驱动 Total Critical Tickets。", "KPI、列表"),
    ("approvalDecision", "派生", "approved/unapproved/other。", "Python 和前端用来过滤 approved rows、repair cost distribution。", "Repair Analytics"),
    ("approvedAmount / unapprovedAmount", "派生", "approved/unapproved 票的金额合计。", "Approved Cost、Repair Analytics、导出。", "成本 KPI"),
    ("avgCostPerTicket", "派生", "Approved Amount / Approved Tickets。", "Current Avg Cost/Ticket KPI 和日期趋势。", "Avg Cost Trend"),
    ("ticketVolumeTrendDaily/Weekly/Monthly", "Python 输出", "按日/周/月的 newTickets、approvedUnapprovedTickets、criticalTotal。", "Critical Status Trends by Date。", "主趋势图"),
    ("daily/weekly/monthlyAvgCostTrend", "Python 输出", "累计 approvedTickets、approvedAmount、avgCostPerTicket。", "Current Avg Cost/Ticket 随日期变化。", "右侧趋势卡"),
    ("teamAssignRows", "Python 输出", "员工当前 critical、removed、avgDaily、avgWeekly。", "Employee Workload 左侧员工工作量。", "Employee Workload"),
    ("approvalRows", "Python 输出/兼容", "旧版按员工 approved/unapproved 汇总。", "当前 index.html 不再使用它驱动右侧员工 approved/unapproved，保留兼容。", "兼容字段"),
    ("repairCostDistribution", "Python 输出", "Approved tickets 按金额区间汇总。", "Tickets Repair Cost Distribution。", "Repair Analytics"),
    ("periodSnapshots", "Python 输出", "任意日期范围预计算 summary 和分布。", "KPI、Repair Analytics、日期范围筛选。", "Dashboard"),
]


def build():
    OUT.parent.mkdir(exist_ok=True)
    doc = Document()
    style_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("C4C / SAP 抓取与分析变量说明")
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("覆盖 fetch_all_tickets_fast_with_firebase_MANDT800_REJECTION_FILTER.py、ctm_v44_history_safe_mandt800_rejection_filter.py 和 index.html 中使用的字段")
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string("555555")

    doc.add_heading("阅读说明", level=1)
    add_bullets(doc, [
        "C4C 字段主要来自 C4C Ticket OData 和 involved-party role API；SAP 字段主要来自 HANA 表 VBAK、VBAP、LIPS、LIKP。",
        "同一个业务字段在代码里可能有多个别名，例如 CreatedOn / createdOn / CreatedAt。文档把这些别名放在同一行。",
        "网页端显示不等于字段没有用途：有些字段只用于判断、筛选、排查或向 Firebase 回写。",
        "Approved 当前团队口径：C4C 状态属于 Z9/Y0/Y1/Y2/Y4/YB，且存在有效 7 开头 10 位 ERPPurchaseOrder，并且订单未 rejected/partially rejected。",
        "Unapproved 当前口径：C4C 状态代码 Y8，或状态文本/历史事件文本包含 unapproved。",
    ])

    add_table(doc, "1. 接口、连接和存储配置变量", CONFIG_FIELDS)
    add_table(doc, "2. C4C 工单核心字段", C4C_CORE)
    add_table(doc, "3. C4C 经销商、维修商与员工字段", C4C_DEALER_EMP)
    add_table(doc, "4. C4C Involved Party / Role 字段", C4C_ROLES)
    add_table(doc, "5. SAP / HANA 订单、物料、交货字段", SAP_FIELDS)
    add_table(doc, "6. Approved / Unapproved 判定相关字段", APPROVAL_FIELDS)
    add_table(doc, "7. 历史事件与状态变化字段", HISTORY_FIELDS)
    add_table(doc, "8. 派生字段、Firebase 分析字段和网页指标字段", DERIVED_FIELDS)

    doc.add_heading("关键业务规则摘要", level=1)
    rules = [
        ("是否已订货/采购", "ERPPurchaseOrder 是采购订单号。代码要求它匹配 7 开头 10 位数字，且 C4C 状态在 approved 状态集合内，才算 approved。"),
        ("采购订单日期", "优先用 PO Document Date / PODocumentDate / DocumentDate；没有时用 ERPPurchaseOrderFirstSeenDate 或 approvalDecisionDate 兜底。"),
        ("未批准关闭", "Y8 或状态文本包含 unapproved 会被归为 unapproved。历史事件里的 toCode/toStatus 也用于抓 unapproved exit。"),
        ("员工归属", "C4C role 40 的 InvolvedPartyName 是员工 Assign To 的权威来源。票面 Assigned To 只作为兜底，Queue Warranty 也保留为可见队列。"),
        ("经销商归属", "优先 DealerName / WarrantyHandlingDealerName；为空时用 role 1001；再经过 dealer alias 表标准化。"),
        ("Parts Delivery", "用 ERPFreeOrder 查 SAP Sales Order，再取 VBAP 物料行和 LIPS/LIKP 交货数。VBAP.ABGRU 有值的 rejected item 会从 Parts Delivery 明细里排除。"),
        ("Critical Status Trends", "蓝/绿线使用 ticketVolumeTrend 的累计 newTickets 和累计 approvedUnapprovedTickets；橙线 criticalTotal 用右侧副坐标。"),
        ("Employee Approved vs Unapproved", "现在直接复用 Repair Analytics 的 repairAnalyticsExportRows() 明细，再按 employee 分组，保证逻辑统一。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "规则", bold=True, color="0B2545")
    set_cell_text(table.rows[0].cells[1], "说明", bold=True, color="0B2545")
    set_cell_shading(table.rows[0].cells[0], "E8EEF5")
    set_cell_shading(table.rows[0].cells[1], "E8EEF5")
    for k, v in rules:
        cells = table.add_row().cells
        set_cell_text(cells[0], k, size=8.5)
        set_cell_text(cells[1], v, size=8.5)
    set_table_geometry(table, (2100, 7260))

    doc.add_heading("代码文件覆盖范围", level=1)
    add_bullets(doc, [
        "fetch_all_tickets_fast_with_firebase_MANDT800_REJECTION_FILTER.py：C4C API 拉数、SAP HANA 查询、Firebase ticket 字段回写、role 数据上传。",
        "ctm_v44_history_safe_mandt800_rejection_filter.py：状态历史、team/dealer/employee analytics、approved/unapproved、avg cost、ticket volume trend 等预计算。",
        "index.html：网页端日期筛选、Repair Analytics 明细、Employee Workload、Critical Status Trends、导出逻辑。",
    ])

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
