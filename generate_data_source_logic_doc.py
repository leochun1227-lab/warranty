from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import sys


DOCS_SKILL_SCRIPTS = Path(
    r"C:\Users\Leo.Li\.codex\plugins\cache\openai-primary-runtime\documents\26.630.12135\skills\documents\scripts"
)
if str(DOCS_SKILL_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(DOCS_SKILL_SCRIPTS))

from table_geometry import apply_table_geometry  # type: ignore


OUT = Path("outputs/parts_delivery_flow_and_repairs_data_source_logic.docx")


PAGE_WIDTH = 6.5
PAGE_WIDTH_DXA = 9360
MARGIN_IN = 1.0

BODY_FONT = "Calibri"
TITLE_SIZE = 20
H1_SIZE = 16
H2_SIZE = 13
H3_SIZE = 12
BODY_SIZE = 11
TABLE_SIZE = 9.5

BLUE = RGBColor.from_string("2E74B5")
DARK_BLUE = RGBColor.from_string("1F4D78")
INK = RGBColor.from_string("0B2545")
TEXT = RGBColor.from_string("111111")
MUTED = RGBColor.from_string("5B6575")
CALL_OUT = RGBColor.from_string("F4F6F9")
TABLE_FILL = RGBColor.from_string("E8EEF5")
SOFT_FILL = RGBColor.from_string("F2F4F7")
LIGHT_BLUE = RGBColor.from_string("F6F9FD")
BORDER = RGBColor.from_string("D9E1EC")
ACCENT_GREEN = RGBColor.from_string("17855A")
ACCENT_ORANGE = RGBColor.from_string("C46B18")
ACCENT_RED = RGBColor.from_string("B0322E")


def set_run_font(run, *, size=None, color=None, bold=None, italic=None, name=BODY_FONT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line


def add_text(doc: Document, text: str, *, size=BODY_SIZE, color=TEXT, bold=False, italic=False, align=None, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_paragraph_spacing(p, after=after)
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def add_title_block(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=3, line=1.0)
    r = p.add_run("Parts Delivery Flow and Repairs Data Source Logic")
    set_run_font(r, size=TITLE_SIZE, color=TEXT, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=8, line=1.0)
    r = p.add_run("面向 CEO 的可读版说明，逐项回答“这个数从哪里来、怎么算出来、什么时候会兜底”。")
    set_run_font(r, size=11.5, color=MUTED)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_text(cell, text, *, size=TABLE_SIZE, bold=False, color=TEXT, align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    set_paragraph_spacing(p, before=0, after=2, line=1.08)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    r = p.add_run(text)
    set_run_font(r, size=size, color=color, bold=bold, italic=italic)
    return p


def set_table_borders(table, color="D9E1EC", size="4"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = tbl_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tbl_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def format_table(table, widths: Sequence[int], *, header_fill="E8EEF5"):
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=sum(widths),
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    set_table_borders(table)
    for idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                set_paragraph_spacing(p, before=0, after=2, line=1.08)
                for run in p.runs:
                    if idx == 0:
                        set_run_font(run, size=TABLE_SIZE, color=INK, bold=True)
                    else:
                        set_run_font(run, size=TABLE_SIZE, color=TEXT)
        if idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, header_fill)


def add_table(doc: Document, title: str, headers: Sequence[str], rows: Iterable[Sequence[str]], widths: Sequence[int], *, caption: str | None = None):
    if title:
        add_h2(doc, title)
    if caption:
        add_note(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT
            if isinstance(value, (int, float)) or str(value).strip().endswith("%"):
                align = WD_ALIGN_PARAGRAPH.RIGHT
            set_cell_text(cells[i], str(value), align=align)
    format_table(table, widths)
    doc.add_paragraph()
    return table


def add_h1(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=18, after=8, line=1.08)
    r = p.add_run(text)
    set_run_font(r, size=H1_SIZE, color=BLUE, bold=True)
    return p


def add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=14, after=6, line=1.08)
    r = p.add_run(text)
    set_run_font(r, size=H2_SIZE, color=BLUE, bold=True)
    return p


def add_h3(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=10, after=4, line=1.08)
    r = p.add_run(text)
    set_run_font(r, size=H3_SIZE, color=DARK_BLUE, bold=True)
    return p


def add_note(doc: Document, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, before=0, after=2, line=1.18)
    r = p.add_run(text)
    set_run_font(r, size=10, color=MUTED)
    apply_table_geometry(
        table,
        [9360],
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    set_table_borders(table, color="D9E1EC")
    doc.add_paragraph()


def add_source_pipeline_section(doc: Document):
    add_h1(doc, "1. 数据管道总览")
    add_note(
        doc,
        "先记住三个层次：页面直接读的是 Firebase 或本地分析包；Firebase 里的历史快照和明细，分别由抓取脚本和聚合脚本算出来；真正决定“为什么是这个值”的，是页面里的兜底顺序和分组规则。",
    )
    add_table(
        doc,
        "",
        ["层级", "文件 / 节点", "页面用途", "一句话解释"],
        [
            ["历史快照", "Firebase: `c4cTickets_test/deliveryFlowHistory/daily.json`", "Parts Delivery Flow 全局 KPI、趋势、总成本", "由 `delivery_flow_aggregator.py` 把每日 snapshot 算好后写入。"],
            ["当前票据明细", "Firebase: `c4cTickets_test/tickets.json`", "Parts Selected / Issued / Open、导出明细、当前 ageing", "由 `fetch_all_tickets_fast_with_firebase_MANDT800_REJECTION_FILTER.py` 写入，包含 C4C + SAP HANA 的票据与 SAP item 明细。"],
            ["成本侧车", "Firebase: `ctmTicketStatusMonitorV44/analytics/deliveryFlow/hanaCostYtd/latest`", "Total Parts Cost 的月度趋势和 PO / fallback 拆分", "若存在，会覆盖票据侧的简化 cost 值。"],
            ["Repair 分析包", "`outputs/repairers_2026/repairers_2026_data.js`", "Repairs 页的 Overview、Top 20、State ranking", "由 `build_repairers_2026_workbook.mjs` 读取 JSON 后生成。"],
            ["Repair drilldown", "`outputs/repairers_2026/repairers_2026_data.js` 或 Firebase tickets fallback", "Repair Insights 详情页", "优先用本地 detail rows，没有时才回 Firebase。"],
            ["地图底图", "`assets/geo/australia-states.geojson` / `assets/geo/new-zealand.geojson`", "Repairs 页的州地图", "只负责形状，不负责业务数值。"],
        ],
        [1300, 2350, 2900, 2810],
    )


def add_glossary_section(doc: Document):
    add_h2(doc, "先说清几个术语")
    add_table(
        doc,
        "",
        ["术语", "业务含义", "代码里怎么判断"],
        [
            ["当前票据（current tickets）", "页面即时拉回的票据明细，带 SAP item 和交货信息。", "来自 `c4cTickets_test/tickets.json`，页面优先用它来重算可见值。"],
            ["快照（snapshot）", "某天系统已经算好的汇总结果。", "来自 `deliveryFlowHistory/daily.json` 里的 `latest` 历史记录。"],
            ["Unique SO", "一个 Sales Order 只算一次。", "按 sales order 去重，不按 ticket 行数去重。"],
            ["Material row", "一个 SAP Sales Order item。", "来自 `Sales Order Details` 数组。"],
            ["Preferred parts cost", "优先按 3110 PO 价格，没有再用 3090 SO 价格。", "CNY 还会按固定 5:1 换算成 AUD。"],
            ["Repair shop ID", "维修商的稳定识别键。", "优先 role 43；没有 role 43 时用 repairer name + state 组合。"],
        ],
        [1500, 3100, 4760],
    )


def parts_delivery_rows():
    return [
        ["顶部 Latest", "页面最上方的最新日期胶囊", "页面先读 `state.history[state.history.length - 1].asOf`，失败时显示 unavailable。", "这是 `deliveryFlowHistory/daily.json` 的最新 snapshot 日期，不是浏览器当前时间。"],
        ["KPI 1: Current Awaiting Parts", "一张蓝色 KPI 卡", "优先用 `awaitingTicketRows(currentTickets).length`。它会按 sales order 去重，只保留 `ticketStatus == YA` 或文本 `awaiting parts`。否则回退到 `latest.awaitingParts.current`，再回退到 `statusMix.awaitingItems` 或 `kpi.openSoItems`。", "这里问的是“有多少个 SO 还在等件”，不是 ticket 行数。"],
        ["KPI 2: Open Parts (Not Issued)", "一张橙色 KPI 卡", "优先用当前票据里的 `materialStats.openCount`。计算规则是：有 material、未 rejected、没有 rejection reason、没有 firstIssueDate。否则回退到 `latest.kpi.openNotIssuedParts`，再回退到 `latest.kpi.openSoItems`。", "这个值是 material row 级别的，不是 ticket 级别。"],
        ["KPI 3: Average Parts Issuing Time", "一张紫色 KPI 卡", "优先用当前票据重新算：`(First Issue Date - SO Created Date)` 的平均天数，只对已发出的 material rows 计入。否则回退到 `latest.kpi.avgDaysToFirstIssue`。", "显示的是“件”的平均首发货时长，不是订单平均时长。"],
        ["Open Parts ageing", "0-7 / 8-14 / 15-30 / 30+ 四条条形", "当前票据存在时，统计所有 open material rows 的 age = `asOf - SO Created Date`。没有当前票据时，直接用 `latest.agingBuckets`。", "这张图回答的是“未发出的件老了多久”。"],
        ["Parts Selected 表格", "两列表格：Period / Sales Order Selected", "当前票据存在时，用 `buildUniqueSeriesFromTickets` 按周/月/年统计 unique SO；没有当前票据时，用历史快照 `flow.newSoItems`。", "表格里的数是 unique SO count，不是 item count。"],
        ["Granularity", "Week / Month / Year 切换", "只是 UI 级别控制，决定 Parts Selected 表格与趋势的聚合粒度。", "不会改变源数据，只改变分组方式。"],
        ["Parts Issued 柱状图", "绿色柱状图", "当前票据存在时，统计 `firstIssueDate` 所在月份的 issued material rows；没有当前票据时，退回 snapshot 的月度序列。", "这里的“issued”是 item 级别的。"],
        ["Average Parts Issuing Time 折线/柱", "蓝柱 + 紫色 3M moving average", "直接读 `latest.avgIssuingTime.byMonth` 和 `latest.avgIssuingTime.volumeByMonth`。", "这是历史聚合结果，不在前端重算。"],
        ["Total Parts Cost Summary", "3 张卡：Total Cost / Avg Parts Cost / Ticket / Nishi E03 Cost", "Total Cost 优先读 `costReport.totalCost`，再回退 `latest.partsCost.total`；Avg 票均价用当前票据数量或 `latest.partsCost.soCount` 作分母；Nishi E03 则读 `latest.nishiE03.ytdCost` 和 `count`。", "Nishi E03 是单独的 YTD 侧车，不和一般 parts cost 混在一起。"],
        ["Total Parts Cost 趋势图", "月度趋势线 + 3110 PO 柱", "优先用 `costReport.months`；没有时回退到 `latest.partsCost.byMonth` / `byMonthPo` / `byMonthFallback`。线代表总 parts cost，柱代表 3110 PO cost。", "页面只显示一个 legend，但底层实际叠了线和柱。"],
        ["Selected row", "中间那条摘要", "只展示 `Total Parts Price`、总金额和 ticket 数。", "这是给人一眼看懂的汇总，不是额外口径。"],
        ["Export Buttons", "3 个导出按钮 + 2 个 KPI 小按钮", "分别导出 Parts Selected、Parts Issued、Total Parts Cost，以及 Current Awaiting / Open Parts 的 ticket scope 导出。", "导出的 Excel 会带 Summary、Trend 和 Ticket Detail 等 sheet。"],
    ]


def add_parts_delivery_section(doc: Document):
    add_h1(doc, "2. Parts Delivery Flow")
    add_note(
        doc,
        "这个页面有两个数据模式。第一种是“快照模式”，直接读历史汇总；第二种是“当前票据模式”，页面一旦拿到 `tickets.json` 就会优先用明细重算可见值，因为这样更接近真实业务现状。",
    )
    add_table(
        doc,
        "",
        ["页面元素", "页面上看到什么", "来源字段 / 节点", "判断逻辑 / 兜底"],
        parts_delivery_rows(),
        [1320, 2140, 3200, 2700],
    )
    add_h3(doc, "这页最容易被问的两句话")
    add_table(
        doc,
        "",
        ["问题", "CEO 版答案"],
        [
            ["为什么有些数是 SO，有些数是 ticket？", "因为页面的业务粒度不是单一的。Parts Selected、Awaiting Parts 以 unique SO 为主；Issued、Open、Ageing、Cost 里很多是 material row 级别。"],
            ["为什么同一页有时候看起来像实时，有时候像历史？", "页面会优先用当前票据明细重算；如果当前票据不可用，才回退到 Firebase 历史快照。这样既能看“今天”，也能保留“昨天怎么算”的痕迹。"],
        ],
        [2200, 7160],
    )


def repairs_overview_rows():
    return [
        ["Overview 标题", "Australia / New Zealand 的地图总览", "读 `workingStates`、`workingWeekly`、`weeklyAnalysis.summary`，并用 `overviewDateRangeLabel()` 算出范围。", "只统计 2026+ 的 repair data；范围来自本地 detail rows 或 tickets fallback 的最早 / 最晚 CreatedOn。"],
        ["Rank metric 按钮", "Price / Tickets / Dealers", "切换 `weeklyMetric`，分别对应 `confirmed_cost`、`ticket_count`、`unique_repairers`。", "它只改变排序和着色，不改数据来源。"],
        ["地图颜色", "州地图上的深浅", "每个州的值来自 `weekStateMetric(row)`，再用 `metricColorScale(value, max)` 显示。", "颜色越深，值越大。"],
        ["地图 tooltip", "Cost / Avg ticket price / Tickets / Dealers", "悬停州时，从 `workingStates` 里取 `confirmed_cost`、`avg_confirmed_cost`、`ticket_count`、`unique_repairers`。", "这个 tooltip 是“州级汇总”，不是 repair shop 级别。"],
        ["州列表表格", "# / State / Price / Tickets / Dealers / Avg Price", "还是 `workingStates`，但经过 `sortAustraliaStates()` 排序。", "默认按当前 metric 排序，平手时按州固定顺序。"],
        ["Top 20 Repairers", "左侧列表", "读 `workingRepairers`，按 `topRepairSort` 排序。", "排序可以切到 ticket_count / confirmed_cost / avg_confirmed_cost。"],
        ["州列表里的 Rank summary", "Rank 区块和表格", "目前 JS 没有单独填 `weekRankSummary` 的卡片内容，真正的数据表就是 `weekStateList`。", "这是一个当前空着的占位容器，不影响主数据。"],
    ]


def repairs_shop_rows():
    return [
        ["Total Cost / Tickets / Avg / Shops", "页面顶部 4 张 KPI 卡", "从 `filtered` 维修商列表汇总：总成本、票数、平均单票成本、维修商数。", "先套 `selectedRepairId` 和搜索条件，再做总和。"],
        ["Shop Summary 表", "维修商清单表", "列是 `repairName`, `uniqueChassisRatio`, `totalCost`, `avgCost`, `ticketCount`, `openCases`, `closedCases`, `chassisTicketCount`, `uniqueChassisCount`。", "`analyzeRepairs()` 先按 repairId 分组，再算票均价、开放 / 已结案、车架重复率。"],
        ["Open cases", "Shop Summary 里的 open cases 列", "当 `invoice_number` 为空，且 `invoice_status` 不是 invoiced / closed 时计为 open。", "这是发票口径，不是工单状态口径。"],
        ["Closed cases", "Shop Summary 里的 closed cases 列", "只要有 invoice number，或者 invoice_status 是 invoiced / closed，就算 closed。", "同一条票不可能同时 open 和 closed。"],
        ["Cost ranges", "右侧 Cost Range Breakdown", "由 `row.costRanges.low/medium/high` 汇总而来。", "规则是 `<500`, `500-2000`, `>=2000`。"],
        ["Repairers dropdown", "隐藏的 Repair shop 筛选", "下拉选项来自 `repairs` 数组。", "主要是为了内部筛选和导出，不是主展示逻辑。"],
        ["Row-level Export", "每行的 Export 按钮", "导出当前 repairId 的 Summary / Repairer Summary / Repairer Tickets。", "只导 2026+ 的票据。"],
        ["Advanced view", "每行的 Advanced view 按钮", "打开 `repair-insights.html?repairId=...`。", "这是钻取页，不是主列表页。"],
    ]


def repairs_repeat_rows():
    return [
        ["Repeated claims on the same chassis", "重复车架面板", "先用 `normalizedChassisFromRaw()` 找车架号，再按 1 次 / 2-3 次 / 4+ 次分桶。", "只看 2026+ tickets，且必须有 repair shop。"],
        ["Summary chips", "Unique chassis / Chassis with repeats / Heavy repeat / Tickets matched", "读 `workingChassis.totals`。", "这些是按车架聚合后的总览，不是维修商总览。"],
        ["Stacked bars", "每个 repeat bucket 的堆叠条", "横向再按 cost buckets 切成 low / mid / high / premium。", "点击某一段会导出该格子的明细。"],
        ["Cost bucket legend", "Low / Medium / High / Premium 标签", "同样来自 `workingChassis.costBuckets`。", "bucket 名称固定，金额阈值固定。"],
        ["Chassis export", "Chassis repeat 的导出 Excel", "导出 Overview / Buckets / Chassis / Tickets 四个 sheet。", "能直接给业务查哪块车架重复、重复了几次、金额是多少。"],
    ]


def add_repairs_section(doc: Document):
    add_h1(doc, "3. Repairs")
    add_note(
        doc,
        "这个页面有两个数据层。上半部分的 Overview 和 Top 20 Repairers 来自本地预计算分析包；下半部分的 Shop Summary 和车架重复分析，会优先用同一份分析包里的 detail rows，没有时才回 Firebase tickets。换句话说，这页不是现场临时拼数，而是“先算好，再展示”。",
    )
    add_table(
        doc,
        "",
        ["页面元素", "页面上看到什么", "来源字段 / 文件", "判断逻辑 / 兜底"],
        repairs_overview_rows(),
        [1380, 2050, 3080, 2850],
    )
    add_table(
        doc,
        "",
        ["页面元素", "页面上看到什么", "来源字段 / 文件", "判断逻辑 / 兜底"],
        repairs_shop_rows() + repairs_repeat_rows(),
        [1380, 2050, 3080, 2850],
    )
    add_h3(doc, "这一页最关键的判断口径")
    add_table(
        doc,
        "",
        ["口径", "解释"],
        [
            ["repair shop 是谁", "优先看 role 43 的 `InvolvedPartyName`。如果没有 role 43，就用 `RepairerName` / `Service Technician` / `DealerName` 再加 state 拼成稳定 ID。"],
            ["成本怎么来", "`repairCostFromTicket()` 依次找 `confirmed_cost_aud`, `confirmed_cost`, `invoice_amount_aud`, `invoice_amount`, `po_amount_aud`, `po_amount`。找到第一个就用它。"],
            ["为什么只看 2026+", "页面硬编码了 `REPAIR_SCOPE_START = 2026-01-01`，所以 2025 的历史不进当前 repair 面板。"],
        ],
        [1900, 7460],
    )


def add_repair_insights_section(doc: Document):
    add_h1(doc, "4. Repair Insights drilldown")
    add_note(
        doc,
        "这是 Repairs 页里点 Advanced view 打开的单店详情页。它不是新的数据源，而是把同一份 detail rows 按一个 `repairId` 再切一层，给单店追问用。",
    )
    add_table(
        doc,
        "",
        ["页面元素", "页面上看到什么", "来源字段 / 文件", "判断逻辑 / 兜底"],
        [
            ["标题 + Repair ID", "店名和 `Repair ID xxx - 2026+ only`", "先从 `repairInfo(firstRow).name` 拿店名，再用查询参数 `repairId`。", "单店 drilldown 的唯一入口就是 query string。"],
            ["Amount chart", "ERP PO / Invoice Cost Distribution", "`repairCostFromTicket()` 的金额按 0-500 / 500-2k / 2k-5k / 5k+ 分桶。", "条形图展示的是票据金额分布，不是维修商金额分布。"],
            ["Status chart", "Status Volume", "按 `TicketStatusText`，没有时退回 `TicketStatus`。", "只取前 8 个状态。"],
            ["Case split pie", "Open vs Closed Cases", "Closed = 有 invoice number 或 invoice_status 为 invoiced/closed。其余为 open。", "这个定义和 Repairs 页的 open/closed 一致。"],
            ["CreatedOn trend", "Ticket CreatedOn Trend", "按 `CreatedOn` 的月份计数。", "页面只要能 parse 日期，就会进月度趋势。"],
            ["Cost by ticket type", "Cost by Ticket Type", "按 `TicketTypeText` 分组后，把成本相加。", "展示的是“花了多少钱”，不是“有多少单”。"],
            ["Repeat cards", "Repeated claims on same chassis", "按 chassis repeat buckets 统计每个 bucket 的车架数、票数和总成本。", "点击卡片可以打开 repeat modal。"],
            ["Repeat modal table", "Chassis / Ticket / Created / Cost / Status / Serial / Dealer", "来自 `buildRepeatState()` 的 `bucket.tickets`。", "这是最细的票据明细层。"],
            ["Export Excel", "Summary / Repeat Ranges / Chassis / Tickets", "当前 repairId 的全部 drilldown 结果。", "导出的是单店明细，不是全局报表。"],
        ],
        [1380, 2050, 3080, 2850],
    )


def add_script_section(doc: Document):
    add_h1(doc, "5. 这些值到底是哪里算出来的")
    add_table(
        doc,
        "",
        ["脚本 / 文件", "职责", "它把什么写进 Firebase 或本地文件"],
        [
            ["`fetch_all_tickets_fast_with_firebase_MANDT800_REJECTION_FILTER.py`", "C4C + SAP HANA 主抓取脚本", "把票据核心字段、roles、Sales Order、SO Created Date、Delivery Count、First Issue Date、Item Rejection Status、Sales Order Details 等写入 `c4cTickets_test/tickets`。"],
            ["`delivery_flow_aggregator.py`", "Parts Delivery Flow 聚合器", "把每日快照写入 `c4cTickets_test/deliveryFlowHistory/daily`，并输出 `partsCost`、`nishiE03`、`costReport`。"],
            ["`build_repairers_2026_workbook.mjs`", "Repair 分析包生成器", "把 `outputs/repairers_2026/repairers_2026_data.json` 转成 `repairers_2026_analysis_state.xlsx` 和页面可读的 `repairers_2026_data.js`。"],
            ["`repair-insights.html`", "单店 drilldown", "不写回数据，只读同一份 detail rows，再按一个 repairId 做二次分析。"],
            ["`build_c4c_sap_variable_dictionary.py`", "字段字典", "把“哪个字段是 C4C，哪个字段是 SAP，哪个字段是派生”的解释统一放在字典里。"],
        ],
        [2240, 2100, 5020],
    )


def build_doc():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", H1_SIZE, BLUE, 18, 10),
        ("Heading 2", H2_SIZE, BLUE, 14, 6),
        ("Heading 3", H3_SIZE, DARK_BLUE, 10, 4),
    ]:
        style = styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.15

    doc.core_properties.title = "Parts Delivery Flow and Repairs Data Source Logic"
    doc.core_properties.subject = "Data source and logic reference"
    doc.core_properties.author = "Codex"

    add_title_block(doc)
    add_note(
        doc,
        "如果某个值前面写的是 `snapshot`，它来自历史汇总；如果写的是 `current tickets`，它来自当前票据明细；如果写的是 `repairers_2026_data.js`，它来自本地预计算分析包。页面看起来像一个仪表盘，但数据层其实分得很清楚。",
    )

    add_source_pipeline_section(doc)
    add_glossary_section(doc)
    add_parts_delivery_section(doc)
    add_repairs_section(doc)
    add_repair_insights_section(doc)
    add_script_section(doc)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    out = build_doc()
    print(out)
